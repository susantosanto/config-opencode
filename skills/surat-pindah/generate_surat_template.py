#!/usr/bin/env python3
"""
Generate Surat dari Template - Custom format untuk surat apapun
Menggunakan template DOCX dengan placeholder {{PLACEHOLDER}}
"""

import requests
import json
import sys
import os
import re
import shutil
from datetime import datetime
from docx import Document

# Default config (bisa di-override via CLI)
DEFAULT_BASE_URL = "http://localhost:5774"
DEFAULT_TOKEN = "AlAiyPRTaYFDKLE"
DEFAULT_NPSN = "20205293"


def get_all_peserta_didik(base_url, token, npsn):
    """Get ALL peserta didik data with pagination"""
    headers = {"Authorization": f"Bearer {token}", "Accept": "application/json"}
    all_data = []
    start = 0
    limit = 200

    while True:
        try:
            params = {"npsn": npsn, "start": start, "limit": limit}
            r = requests.get(
                f"{base_url}/WebService/getPesertaDidik",
                headers=headers,
                params=params,
                timeout=15,
            )
            if r.status_code != 200:
                break
            data = r.json()
            rows = data.get("rows", [])
            if not rows:
                break
            all_data.extend(rows)
            total = data.get("results", 0)
            start += limit
            if start >= total:
                break
        except Exception as e:
            print(f"Error fetching data: {e}")
            break
    return all_data


def search_students(nama_cari, base_url, token, npsn):
    """Cari siswa berdasarkan nama dari Dapodik"""
    semua_data = get_all_peserta_didik(base_url, token, npsn)
    if not semua_data:
        print("Gagal mengambil data dari Dapodik. Pastikan Dapodik berjalan.")
        return []

    hasil = []
    nama_cari_lower = nama_cari.lower()
    for siswa in semua_data:
        nama_siswa = siswa.get("nama", "").lower()
        if nama_cari_lower in nama_siswa or nama_siswa in nama_cari_lower:
            hasil.append(siswa)
    return hasil


def get_student_data_dapodik(
    nama_cari, base_url, token, npsn, auto_select=False, choice_idx=None
):
    """Cari dan pilih siswa dari Dapodik, return data dict"""
    students = search_students(nama_cari, base_url, token, npsn)
    if not students:
        print(f"[ERROR] Tidak ditemukan siswa dengan nama '{nama_cari}'")
        return None

    if len(students) > 1:
        print("\n" + "=" * 70)
        print("DITEMUKAN BEBERAPA SISWA:")
        print("=" * 70)
        for i, s in enumerate(students, 1):
            print(f"\n{i}. {s.get('nama', '-')}")
            print(f"   NISN: {s.get('nisn', '-')}")
            print(f"   Kelas: {s.get('nama_rombel', '-')}")
            print(
                f"   TTL: {s.get('tempat_lahir', '-')}, {s.get('tanggal_lahir', '-')}"
            )

        if auto_select:
            if choice_idx is not None and 0 <= choice_idx < len(students):
                return students[choice_idx]
            return students[0]

        try:
            choice = input("\nPilih nomor siswa (1-{}): ".format(len(students)))
            if choice.strip() == "0":
                return None
            idx = int(choice) - 1
            if 0 <= idx < len(students):
                return students[idx]
        except (ValueError, EOFError):
            return students[0]  # fallback auto-select

    return students[0]


def build_student_data_dict(siswa):
    """Convert Dapodik student data to flat dict for template placeholders"""
    import re

    nama_rombel = siswa.get("nama_rombel", "-")
    jenis_kelamin = (
        "Laki-laki"
        if siswa.get("jenis_kelamin") == "L"
        else "Perempuan"
        if siswa.get("jenis_kelamin") == "P"
        else "-"
    )

    # Kelas format: III (Tiga)
    tingkat_map = {
        "1": ("I", "Satu"),
        "2": ("II", "Dua"),
        "3": ("III", "Tiga"),
        "4": ("IV", "Empat"),
        "5": ("V", "Lima"),
        "6": ("VI", "Enam"),
    }
    kelas_str = "-"
    kelas_romawi = "-"
    match = re.search(r"(\d+)", str(nama_rombel).lower())
    if match and match.group(1) in tingkat_map:
        romawi, kata = tingkat_map[match.group(1)]
        kelas_str = f"{romawi} ({kata})"
        kelas_romawi = romawi

    # Alamat
    dusun = siswa.get("nama_dusun", "")
    alamat_jalan = siswa.get("alamat_jalan", "")
    desa = siswa.get("desa_kelurahan", "")
    kecamatan = siswa.get("kecamatan", "")
    kabupaten = siswa.get("kabupaten_kota", "")
    provinsi = siswa.get("provinsi", "")

    addr_parts_1 = []
    if dusun and dusun != "-":
        addr_parts_1.append(f"Kp. {dusun}")
    elif alamat_jalan and alamat_jalan != "-":
        addr_parts_1.append(alamat_jalan)
    if desa and desa != "-":
        addr_parts_1.append(f"Ds. {desa}")
    if kecamatan and kecamatan != "-":
        addr_parts_1.append(f"Kec. {kecamatan}")
    alamat_line1 = " ".join(addr_parts_1)

    addr_parts_2 = []
    if kabupaten and kabupaten != "-":
        addr_parts_2.append(f"Kab. {kabupaten}")
    if provinsi and provinsi != "-":
        addr_parts_2.append(f"Prov. {provinsi}")
    alamat_line2 = " ".join(addr_parts_2)

    today = datetime.now()
    bulan_map = {
        1: "Januari",
        2: "Februari",
        3: "Maret",
        4: "April",
        5: "Mei",
        6: "Juni",
        7: "Juli",
        8: "Agustus",
        9: "September",
        10: "Oktober",
        11: "November",
        12: "Desember",
    }
    tgl_surat = f"{today.day} {bulan_map[today.month]} {today.year}"
    tahun_surat = today.strftime("%Y")

    return {
        "{{NAMA_SISWA}}": siswa.get("nama", "-"),
        "{{NISN}}": siswa.get("nisn", "-"),
        "{{NIK}}": siswa.get("nik", "-"),
        "{{JENIS_KELAMIN}}": jenis_kelamin,
        "{{KELAS}}": kelas_str,
        "{{KELAS_ROMAWI}}": kelas_romawi,
        "{{KELAS_ANGKA}}": match.group(1) if match else "-",
        "{{NAMA_AYAH}}": siswa.get("nama_ayah", "-")
        if siswa.get("nama_ayah") and siswa.get("nama_ayah") != "-"
        else "-",
        "{{NAMA_IBU}}": siswa.get("nama_ibu", "-")
        if siswa.get("nama_ibu") and siswa.get("nama_ibu") != "-"
        else "-",
        "{{NAMA_ORTU}}": siswa.get("nama_ayah", "-")
        if siswa.get("nama_ayah") and siswa.get("nama_ayah") != "-"
        else "-",
        "{{ALAMAT_LINE1}}": alamat_line1,
        "{{ALAMAT_LINE2}}": alamat_line2,
        "{{ALAMAT_LENGKAP}}": f"{alamat_line1} {alamat_line2}".strip(),
        "{{TEMPAT_LAHIR}}": siswa.get("tempat_lahir", "-"),
        "{{TANGGAL_LAHIR}}": siswa.get("tanggal_lahir", "-"),
        "{{AGAMA}}": siswa.get("agama", "-"),
        "{{NO_KK}}": siswa.get("no_kk", "-"),
        "{{NAMA_ROMBEL}}": nama_rombel,
        "{{TGL_SURAT}}": tgl_surat,
        "{{TAHUN_SURAT}}": tahun_surat,
        "{{DESA}}": desa,
        "{{KECAMATAN}}": kecamatan,
        "{{KABUPATEN}}": kabupaten,
        "{{PROVINSI}}": provinsi,
    }


def replace_placeholders_in_run(run, replacements):
    """Replace placeholders in a single run's text"""
    original_text = run.text
    new_text = original_text
    for placeholder, value in replacements.items():
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(value))
    if new_text != original_text:
        run.text = new_text
        return True
    return False


def generate_from_template(template_path, output_path, replacements):
    """
    Generate surat dari template dengan replace placeholder.
    Placeholder format: {{NAMA_PLACEHOLDER}}

    Args:
        template_path: Path ke file template DOCX
        output_path: Path ke file output DOCX
        replacements: Dict mapping placeholder -> value
    """
    if not os.path.exists(template_path):
        print(f"[ERROR] Template tidak ditemukan: {template_path}")
        return None

    # Copy template
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    shutil.copy2(template_path, output_path)

    # Open and replace
    doc = Document(output_path)

    replaced_count = 0
    # Replace in paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            if replace_placeholders_in_run(run, replacements):
                replaced_count += 1

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if replace_placeholders_in_run(run, replacements):
                            replaced_count += 1

    # Replace in headers/footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            for run in p.runs:
                replace_placeholders_in_run(run, replacements)
        for p in section.footer.paragraphs:
            for run in p.runs:
                replace_placeholders_in_run(run, replacements)

    doc.save(output_path)

    # Report
    print(f"\n[OK] Surat berhasil dibuat!")
    print(f"[OK] Template: {template_path}")
    print(f"[OK] Output: {output_path}")
    print(f"[OK] Placeholder replaced: {replaced_count}")
    return output_path


def load_custom_replacements(json_path):
    """Load custom replacements from JSON file"""
    if not os.path.exists(json_path):
        print(f"[ERROR] JSON tidak ditemukan: {json_path}")
        return {}
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def print_usage():
    print("""
CARA PENGGUNAAN:

1. Dengan data siswa dari Dapodik:
   python generate_surat_template.py --template "path/template.docx" --output "path/output.docx" --student "nama siswa"

2. Dengan data custom dari JSON:
   python generate_surat_template.py --template "path/template.docx" --output "path/output.docx" --data "path/data.json"

3. Dengan data siswa + custom tambahan:
   python generate_surat_template.py --template "path/template.docx" --output "path/output.docx" --student "nama siswa" --data "path/extra.json"

4. Dengan output folder (auto nama file):
   python generate_surat_template.py --template "path/template.docx" --output-folder "path/folder/" --student "nama siswa"

PLACEHOLDER YANG TERSEDIA:
  {{NAMA_SISWA}}    {{NISN}}           {{NIK}}
  {{JENIS_KELAMIN}} {{KELAS}}          {{KELAS_ROMAWI}}
  {{NAMA_AYAH}}     {{NAMA_IBU}}       {{NAMA_ORTU}}
  {{ALAMAT_LINE1}}  {{ALAMAT_LINE2}}   {{ALAMAT_LENGKAP}}
  {{TEMPAT_LAHIR}}  {{TANGGAL_LAHIR}}  {{AGAMA}}
  {{NO_KK}}         {{NAMA_ROMBEL}}    {{TGL_SURAT}}
  {{TAHUN_SURAT}}   {{DESA}}           {{KECAMATAN}}
  {{KABUPATEN}}     {{PROVINSI}}

  Custom placeholder di JSON: {{CUSTOM_KEY}}
""")


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Generate surat dari template DOCX")
    parser.add_argument("--template", "-t", help="Path ke file template DOCX")
    parser.add_argument("--output", "-o", help="Path ke file output DOCX")
    parser.add_argument("--output-folder", "-of", help="Folder output (auto nama file)")
    parser.add_argument("--student", "-s", help="Nama siswa (cari dari Dapodik)")
    parser.add_argument("--data", "-d", help="Path JSON data custom")
    parser.add_argument("--auto", action="store_true", help="Auto-select siswa pertama")
    parser.add_argument("--choice", type=int, help="Pilih siswa ke-N")
    parser.add_argument("--dapodik-url", default=DEFAULT_BASE_URL, help="URL Dapodik")
    parser.add_argument("--token", default=DEFAULT_TOKEN, help="Token Dapodik")
    parser.add_argument("--npsn", default=DEFAULT_NPSN, help="NPSN sekolah")
    parser.add_argument(
        "--list-placeholders",
        action="store_true",
        help="Tampilkan placeholder tersedia",
    )

    args = parser.parse_args()

    if args.list_placeholders:
        print_usage()
        return

    if not args.template:
        print("[ERROR] Template harus diisi (--template)")
        print_usage()
        sys.exit(1)

    if not args.output and not args.output_folder:
        print("[ERROR] Output harus diisi (--output atau --output-folder)")
        print_usage()
        sys.exit(1)

    # Build replacements
    replacements = {}

    # From Dapodik student
    if args.student:
        siswa = get_student_data_dapodik(
            args.student,
            args.dapodik_url,
            args.token,
            args.npsn,
            auto_select=args.auto,
            choice_idx=args.choice - 1 if args.choice else None,
        )
        if not siswa:
            sys.exit(1)
        replacements = build_student_data_dict(siswa)
        print(f"\nData siswa: {siswa.get('nama')} (NISN: {siswa.get('nisn')})")

    # From custom JSON
    if args.data:
        custom = load_custom_replacements(args.data)
        # Convert keys to placeholder format
        for key, value in custom.items():
            placeholder = key if key.startswith("{{") else f"{{{{{key}}}}}"
            replacements[placeholder] = value

    if not replacements:
        print("[ERROR] Tidak ada data. Gunakan --student atau --data")
        print_usage()
        sys.exit(1)

    # Build output path
    if args.output:
        output_path = args.output
    else:
        # Auto name from student or template
        nama_siswa = replacements.get("{{NAMA_SISWA}}", "document")
        nama_file = nama_siswa.lower().replace(" ", "_")
        template_name = os.path.splitext(os.path.basename(args.template))[0]
        output_path = os.path.join(
            args.output_folder, f"{template_name}_{nama_file}.docx"
        )

    # Generate
    generate_from_template(args.template, output_path, replacements)


if __name__ == "__main__":
    main()

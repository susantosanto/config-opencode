#!/usr/bin/env python3
"""
Generate Surat Pindah Sekolah - Otomatis dari data Dapodik
Menggunakan template asli SK Pindah Sekolah.docx
"""

import requests
import json
import sys
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_URL = "http://localhost:5774"
TOKEN = "AlAiyPRTaYFDKLE"
NPSN = "20205293"
SEKOLAH = "SD NEGERI PASIRHALANG"
ALAMAT_SEKOLAH = "Kp. Pasirhalang RT.03 RW.14"
DESA = "Mandalamukti"
KECAMATAN = "Cikalongwetan"
KABUPATEN = "Bandung Barat"
PROVINSI = "Jawa Barat"

headers = {"Authorization": f"Bearer {TOKEN}", "Accept": "application/json"}


def get_all_peserta_didik():
    """Get ALL peserta didik data with pagination"""
    all_data = []
    start = 0
    limit = 200

    while True:
        try:
            params = {"npsn": NPSN, "start": start, "limit": limit}
            r = requests.get(
                f"{BASE_URL}/WebService/getPesertaDidik",
                headers=headers,
                params=params,
                timeout=15,
            )

            if r.status_code != 200:
                break

            data = r.json()
            rows = data.get("rows", [])
            if not rows:
                break

            all_data.extend(rows)
            total = data.get("results", 0)
            start += limit
            if start >= total:
                break
        except Exception as e:
            print(f"Error fetching data: {e}")
            break

    return all_data


def search_students(nama_cari):
    """Cari siswa berdasarkan nama"""
    semua_data = get_all_peserta_didik()

    if not semua_data:
        print("Gagal mengambil data dari Dapodik. Pastikan Dapodik berjalan.")
        return []

    hasil = []
    nama_cari_lower = nama_cari.lower()

    for siswa in semua_data:
        nama_siswa = siswa.get("nama", "").lower()
        if nama_cari_lower in nama_siswa or nama_siswa in nama_cari_lower:
            hasil.append(siswa)

    return hasil


def display_student_options(students):
    """Tampilkan daftar siswa untuk dipilih"""
    print("\n" + "=" * 70)
    print("DITEMUKAN BEBERAPA SISWA:")
    print("=" * 70)

    for i, siswa in enumerate(students, 1):
        print(f"\n{i}. {siswa.get('nama', '-')}")
        print(f"   NISN: {siswa.get('nisn', '-')}")
        print(f"   Kelas: {siswa.get('nama_rombel', '-')}")
        print(
            f"   TTL: {siswa.get('tempat_lahir', '-')}, {siswa.get('tanggal_lahir', '-')}"
        )
        print(f"   Nama Ayah: {siswa.get('nama_ayah', '-')}")
        print(f"   Nama Ibu: {siswa.get('nama_ibu', '-')}")

    print("\n" + "-" * 70)
    print("0. BATALKAN")
    print("-" * 70)


def select_student(students, non_interactive=False, choice_idx=None):
    """Minta user memilih siswa, atau pilih otomatis jika non-interactive"""
    if non_interactive:
        if choice_idx is not None and 0 <= choice_idx < len(students):
            return students[choice_idx]
        # Auto-select first if no choice given
        return students[0]

    while True:
        try:
            choice = input("\nPilih nomor siswa (1-{}): ".format(len(students)))
            if choice.strip() == "0":
                print("Operasi dibatalkan.")
                return None

            choice_idx = int(choice) - 1
            if 0 <= choice_idx < len(students):
                return students[choice_idx]
            else:
                print(f"Nomor tidak valid. Pilih 1-{len(students)} atau 0 untuk batal.")
        except ValueError:
            print("Input tidak valid. Masukkan angka.")


def get_kelas_tingkat(nama_rombel):
    """Konversi nama rombel ke format kelas dan tingkat"""
    if not nama_rombel:
        return "-", "-"

    rombel_lower = nama_rombel.lower()

    # Extract angka dari nama rombel
    match = re.search(r"(\d+)", rombel_lower)
    if match:
        angka = match.group(1)
        tingkat_map = {"1": "I", "2": "II", "3": "III", "4": "IV", "5": "V", "6": "VI"}
        tingkat_romawi = tingkat_map.get(angka, angka)
        return f"{angka} ({tingkat_romawi})", tingkat_romawi

    return nama_rombel, nama_rombel


def create_surat_pindah(siswa, output_folder, template_path=None):
    """Buat surat pindah sekolah berdasarkan template asli"""

    os.makedirs(output_folder, exist_ok=True)

    nama_siswa = siswa.get("nama", "siswa")
    nama_file = f"surat_mutasi_{nama_siswa.lower().replace(' ', '_')}.docx"
    output_path = os.path.join(output_folder, nama_file)

    # Ambil data siswa
    nisn = siswa.get("nisn", "-")
    nama_ayah = siswa.get("nama_ayah", "-")
    nama_ibu = siswa.get("nama_ibu", "-")
    nama_rombel = siswa.get("nama_rombel", "-")
    jenis_kelamin = (
        "Laki-laki"
        if siswa.get("jenis_kelamin") == "L"
        else "Perempuan"
        if siswa.get("jenis_kelamin") == "P"
        else "-"
    )

    # Ambil alamat dari data siswa atau gunakan default sekolah
    alamat_jalan = siswa.get("alamat_jalan", "")
    rt = siswa.get("rt", "")
    rw = siswa.get("rw", "")
    dusun = siswa.get("nama_dusun", "")
    desa = siswa.get("desa_kelurahan", DESA)
    kecamatan = siswa.get("kecamatan", KECAMATAN)
    kabupaten = siswa.get("kabupaten_kota", KABUPATEN)
    provinsi = siswa.get("provinsi", PROVINSI)

    # Format alamat lengkap
    alamat_parts = []
    if dusun and dusun != "-":
        alamat_parts.append(f"Kp. {dusun}")
    elif alamat_jalan and alamat_jalan != "-":
        alamat_parts.append(alamat_jalan)
    if rt and rt != "-":
        alamat_parts.append(f"RT.{rt}")
    if rw and rw != "-":
        alamat_parts.append(f"RW.{rw}")
    if desa and desa != "-":
        alamat_parts.append(f"Ds. {desa}")
    if kecamatan and kecamatan != "-":
        alamat_parts.append(f"Kec. {kecamatan}")
    if kabupaten and kabupaten != "-":
        alamat_parts.append(f"Kab. {kabupaten}")

    alamat_lengkap = (
        " ".join(alamat_parts)
        if alamat_parts
        else f"Kp. {desa} Kec. {kecamatan} Kab. {kabupaten}"
    )

    # Konversi kelas
    kelas_str, tingkat_romawi = get_kelas_tingkat(nama_rombel)

    # Tanggal surat
    today = datetime.now()
    bulan_map = {
        1: "Januari",
        2: "Februari",
        3: "Maret",
        4: "April",
        5: "Mei",
        6: "Juni",
        7: "Juli",
        8: "Agustus",
        9: "September",
        10: "Oktober",
        11: "November",
        12: "Desember",
    }
    tgl_surat = f"{today.day} {bulan_map[today.month]} {today.year}"
    nomor_surat = f"421.2/{today.strftime('%Y')}/SD - 007/{today.strftime('%Y')}"

    # Buat dokumen baru
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    # --- JUDUL SURAT ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SURAT KETERANGAN PINDAH SEKOLAH")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # --- NOMOR SURAT ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Nomor: {nomor_surat}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Spasi
    doc.add_paragraph()

    # --- PEMBUKA ---
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        "Yang bertanda tangan dibawah ini, Kepala Sekolah Dasar Negeri Pasirhalang Desa Mandalamukti Kecamatan Cikalongwetan Kabupaten Bandung Barat dengan ini menyatakan bahwa :"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- DATA SISWA ---
    data_siswa = [
        ("Nama", nama_siswa),
        ("NISN/ No. Induk", nisn),
        ("Jenis Kelamin", jenis_kelamin),
        ("Murid Kelas", kelas_str),
    ]

    for label, value in data_siswa:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f"{label:<20}: {value}")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    # Spasi
    doc.add_paragraph()

    # --- SURAT PERMOHONAN DARI ORANG TUA ---
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        "Sesuai dengan Surat Keterangan Permohonan Pindah dari orang tua/wali murid :"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    data_ortu = [
        ("Nama", nama_ayah if nama_ayah and nama_ayah != "-" else nama_siswa),
        ("Ibu", nama_ibu if nama_ibu and nama_ibu != "-" else "-"),
        ("Alamat", alamat_lengkap),
    ]

    for label, value in data_ortu:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f"{label:<20}: {value}")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    # Spasi
    doc.add_paragraph()

    # --- TUJUAN PINDAH ---
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        "Telah mengajukan pindah sekolah ke SD Negeri Girimukti di Kecamatan Cikalongwetan Kabupaten Bandung Barat - Jawabarat, bersama ini sertakan Laporan Hasil Belajar Siswa (LAPOR)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Spasi besar untuk tanda tangan
    for _ in range(6):
        doc.add_paragraph()

    # --- TANDA TANGAN ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Mandalamukti, {tgl_surat}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Kepala Sekolah")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Spasi untuk tanda tangan
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("(.........................)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- PEMISAH HALAMAN ---
    doc.add_page_break()

    # --- HALAMAN KEDUA (Untuk diisi sekolah tujuan) ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SURAT KETERANGAN PINDAH SEKOLAH")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Nomor: {nomor_surat}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        "Setelah anak tersebut diterima di sekolah ini, isian dibawah ini harap diisi dan lembar kedua di kirim kembali pada kami."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # --- TABEL DATA SEKOLAH TUJUAN ---
    data_sekolah_tujuan = [
        ("Nama Sekolah", ""),
        ("Status Sekolah", ""),
        ("Alamat", ""),
        ("Desa/ Kelurahan", ""),
        ("Kec/ Kab", ""),
        ("Provinsi", ""),
        ("Diterima Tanggal", ""),
        ("Di Tingkat/ Kelas", ""),
    ]

    for label, value in data_sekolah_tujuan:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)
        dots = "." * 60
        run = p.add_run(f"{label:<20}: {dots}")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    # Spasi
    for _ in range(4):
        doc.add_paragraph()

    # Tanda tangan sekolah tujuan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Kepala Sekolah")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("(.........................)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Simpan
    doc.save(output_path)
    print(f"\n[OK] Surat pindah berhasil dibuat!")
    print(f"[OK] File: {output_path}")
    return output_path


def main():
    print("=" * 70)
    print("GENERATE SURAT PINDAH SEKOLAH")
    print("=" * 70)

    # Check Dapodik
    try:
        r = requests.get(BASE_URL, timeout=5)
        if r.status_code != 200:
            print("[ERROR] Dapodik tidak merespons. Pastikan Dapodik berjalan.")
            sys.exit(1)
    except:
        print("[ERROR] Dapodik tidak dapat diakses di http://localhost:5774")
        sys.exit(1)

    # Parse arguments
    nama_input = None
    choice_idx = None
    non_interactive = False

    if len(sys.argv) > 1:
        nama_input = sys.argv[1]
    if len(sys.argv) > 2:
        arg2 = sys.argv[2]
        if arg2 == "--auto":
            non_interactive = True
        else:
            try:
                choice_idx = int(arg2) - 1
                non_interactive = True
            except ValueError:
                pass
    if len(sys.argv) > 3 and sys.argv[3] == "--auto":
        non_interactive = True

    if not nama_input:
        try:
            nama_input = input("\nMasukkan nama siswa: ").strip()
        except EOFError:
            print("[ERROR] Nama siswa harus diberikan sebagai argumen.")
            sys.exit(1)

    if not nama_input:
        print("[ERROR] Nama siswa harus diisi.")
        sys.exit(1)

    print(f"\nMencari siswa dengan nama: '{nama_input}'...")

    # Search students
    students = search_students(nama_input)

    if not students:
        print(f"\n[TIDAK DITEMUKAN] Tidak ada siswa dengan nama '{nama_input}'")
        print("\nTips:")
        print("  - Coba gunakan nama depan saja")
        print("  - Pastikan ejaan nama benar")
        print("  - Cek apakah Dapodik sudah berjalan")
        sys.exit(1)

    # Multiple results
    if len(students) > 1:
        display_student_options(students)

        # Output JSON for Telegram bot parsing
        result_json = {
            "status": "multiple_found",
            "count": len(students),
            "students": [],
        }
        for i, s in enumerate(students, 1):
            result_json["students"].append(
                {
                    "number": i,
                    "nama": s.get("nama", "-"),
                    "nisn": s.get("nisn", "-"),
                    "kelas": s.get("nama_rombel", "-"),
                    "ttl": f"{s.get('tempat_lahir', '-')}, {s.get('tanggal_lahir', '-')}",
                }
            )

        print(
            f"\n[JSON_OUTPUT]{json.dumps(result_json, ensure_ascii=False)}[/JSON_OUTPUT]"
        )

        if non_interactive and choice_idx is not None:
            selected = select_student(
                students, non_interactive=True, choice_idx=choice_idx
            )
        elif non_interactive:
            # Auto-select first
            selected = students[0]
            print(
                f"\n[INFO] Multiple students found. Auto-selected: {selected.get('nama')}"
            )
        else:
            try:
                selected = select_student(students, non_interactive=False)
            except EOFError:
                # Fallback to auto-select when running non-interactively
                selected = students[0]
                print(
                    f"\n[INFO] Multiple students found. Auto-selected: {selected.get('nama')}"
                )

        if not selected:
            sys.exit(0)
        siswa = selected
    else:
        siswa = students[0]
        print(
            f"\nDitemukan: {siswa.get('nama')} (NISN: {siswa.get('nisn')}) - Kelas: {siswa.get('nama_rombel')}"
        )

        if not non_interactive:
            confirm = input("\nLanjutkan dengan siswa ini? (y/n): ").strip().lower()
            if confirm != "y":
                print("Operasi dibatalkan.")
                sys.exit(0)

    # Generate surat
    output_folder = r"C:\Users\USER\Documents\SK_Pindah"
    output_path = create_surat_pindah(siswa, output_folder)

    # Output result JSON for Telegram bot
    result_json = {
        "status": "success",
        "student": {
            "nama": siswa.get("nama", "-"),
            "nisn": siswa.get("nisn", "-"),
            "kelas": siswa.get("nama_rombel", "-"),
        },
        "file": output_path,
    }
    print(f"\n[JSON_OUTPUT]{json.dumps(result_json, ensure_ascii=False)}[/JSON_OUTPUT]")


if __name__ == "__main__":
    main()
